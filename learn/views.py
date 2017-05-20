from django.contrib.auth.models import User
from django.contrib.auth import get_user_model
from django.contrib.auth import get_user
from django.db.models import Sum
from django.shortcuts import render, render_to_response, get_object_or_404, HttpResponseRedirect
from django.conf import settings
from django import http
from rest_framework import views, viewsets, filters
from rest_framework.reverse import reverse
from rest_framework.response import Response
from rest_framework.views import APIView

from learn.serializers import *
from learn.models import *
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticatedOrReadOnly
from learn.permissions import AdminOrReadOnly, IsOwnerOrReadOnly
from rest_framework import generics
from filters.mixins import FiltersMixin
# from docx import Document
import docx
# from learn.forms import *


# Create your views here.
def index(request):
    context_dict = locals()
    return render(request, "index.html", context_dict)


class UserViewSet(viewsets.ModelViewSet):
    queryset = User.objects.all()
    serializer_class = UserSerializer
    permission_classes = (AdminOrReadOnly,)


class ExamsViewSet(viewsets.ModelViewSet):
    queryset = Exams.objects.all()
    serializer_class = ExamsSerializer
    permission_classes = (AdminOrReadOnly,)


class ExamSubjectViewSet(viewsets.ModelViewSet):
    queryset = ExamSubject.objects.all()
    serializer_class = ExamSubjectSerializer
    permission_classes = (AdminOrReadOnly,)


class ExamList(generics.ListAPIView):
    serializer_class = ExamSubjectSerializer
    permission_classes = (AdminOrReadOnly,)

    def get_queryset(self):
        exam = self.kwargs['exam']
        return ExamSubject.objects.filter(exam=exam)


class QuestionViewSet(FiltersMixin, viewsets.ModelViewSet):
    queryset = Question.objects.all()
    serializer_class = QuestionSerializer
    permission_classes = (AdminOrReadOnly,)


class QuestionList(generics.ListAPIView):
    serializer_class = QuestionSerializer
    permission_classes = (AdminOrReadOnly,)

    def get_queryset(self):
        exam = self.kwargs['exam']
        year = self.kwargs['year']
        subject = self.kwargs['subject']
        return Question.objects.filter(exam__exam=exam, exam__year=year, exam__subject=subject)


class AnswerViewSet(viewsets.ModelViewSet):
    queryset = Answer.objects.all()
    serializer_class = AnswerSerializer
    permission_classes = (AdminOrReadOnly,)


class UserScoreViewSet(viewsets.ModelViewSet):
    queryset = UserScore.objects.all()
    serializer_class = UserScoreSerializer
    permission_classes = (IsOwnerOrReadOnly,)


class LeadersBoardViewSet(viewsets.ModelViewSet):
    queryset = User.objects.annotate(total_points=Sum('myscores__score')).order_by('-total_points')
    serializer_class = LeadersBoardSerializer
    permission_classes = (IsOwnerOrReadOnly,)


# def upload(request):
#     template = 'bulkaddquestions.html'
#     if request.method == 'POST':
#         if 'file' in request.FILES:
#             file = request.FILES['file']
#             filename = file['filename']
#             doc = docx.Document(filename)
#             fulltext = []
#             for para in doc.paragraphs:
#                 fulltext.append(para.text)
#
#             return http.HttpResponseRedirect('upload_success.html')
#     else:
#         form = QuestionsForm()
#         return render_to_response(template, {'form': form})
#
#     return render(request, 'bulkaddquestions.html', {'form': form})
#
#
# # import docx
# def gettext(filename):
#     doc = docx.Document(filename)
#     fulltext = []
#     for para in doc.paragraphs:
#         fulltext.append(para.text)
#     return '\n'.join(fulltext)

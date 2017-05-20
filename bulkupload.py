import docx
from learn.models import *

doc = docx.Document('Government1991.docx')


def get_questions(e_doc):
    indexes = []
    for i, paragraph in enumerate(e_doc.paragraphs):
        # import pdb; pdb.set_trace()
        try:
            if paragraph.text[0] == '#':
                indexes.append(i)
        except IndexError:
            pass
    # import pdb; pdb.set_trace()
    result = []
    for i, j in enumerate(indexes):
        if i + 1 < len(indexes):
            result.append(e_doc.paragraphs[j:indexes[i + 1]])
    result.append(e_doc.paragraphs[indexes[-1]:])
    return result


def get_answers(question):
    answers = []
    for paragraph in question:
        try:
            if paragraph.text[0] == '@':
                answers.append(paragraph)
        except IndexError:
            pass
    valid_answer = [x.text for x in answers if x.text[:2] == '@@']
    if len(valid_answer) > 0:
        valid_answer = valid_answer[0]
    else:
        valid_answer = ''
    answer_text = [x.text for x in answers]
    question_text = [x.text for x in question if x not in answers]
    return answer_text, valid_answer, question_text


class Questions(object):

    def __init__(self, *args):
        self.id = args[0]
        self.question = " ".join(args[3])
        self.answers = [x.replace('@', '') for x in args[1]]
        self.valid_answer = args[2].replace('@@', '')

    def __repr__(self):
        return "<Question {}>".format(self.id)

# all_questions contain each of the questions
all_questions = []
for i, question in enumerate(get_questions(doc)):
    all_questions.append(Questions(i, *get_answers(question)))

# populating the database from all_questions
for i in range(len(all_questions)):
    exam = ExamSubject.objects.get(exam='JAMB', subject='Government', year='1991')
    question_tex = all_questions[i].question
    question_text = question_tex[1:]
    valid = all_questions[i].valid_answer
    answers = all_questions[i].answers
    Question.objects.create(question_text=question_text, serial_no=i, is_published=True, exam=exam, correct=valid)
    question2ans = Question.objects.get(serial_no=i)
    for answer in answers:
        Answer.objects.create(text=answer, question=question2ans)


